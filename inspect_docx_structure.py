"""
Script to inspect the structure of Word documents to understand their format
"""

from docx import Document
import os

base_path = r"C:\Users\tekzo\Downloads\MCQ Question Bank Mathematics"

# Check Grade 10 document as a sample
file_path = os.path.join(base_path, "Question Bank Grade 10 Mathematics MCQ.docx")

if os.path.exists(file_path):
    doc = Document(file_path)
    print(f"\nInspecting: {os.path.basename(file_path)}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("\nFirst 50 paragraphs:")
    print("="*80)
    
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if text:  # Only show non-empty paragraphs
            print(f"{i:3d}: {text[:100]}")
    
    print("\n" + "="*80)
    print("\nChecking for tables...")
    print(f"Total tables: {len(doc.tables)}")
    
    if len(doc.tables) > 0:
        print("\nFirst table structure:")
        table = doc.tables[0]
        print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        if len(table.rows) > 0:
            print("\nFirst row:")
            for j, cell in enumerate(table.rows[0].cells):
                print(f"  Col {j}: {cell.text.strip()[:50]}")
else:
    print(f"File not found: {file_path}")

